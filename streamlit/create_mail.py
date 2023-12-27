import streamlit as st
import google.generativeai as palm
from docx import Document
import base64
import datetime
import random
import string

from io import BytesIO

from googletrans import Translator

translator = Translator()

# Konfigurasi API
palm.configure(api_key="AIzaSyDS__6q4C6Hh3fdaSMpuX_mxAJe-f354J8")

defaults_bot1 = {
    'model': 'models/chat-bison-001',
    'temperature': 0.25,
    'candidate_count': 1,
    'top_k': 40,
    'top_p': 0.95,
}

context_bot1 = "Given a topic, write emails in a concise, professional manner for"
context_bot2 = "Given a topic, write cv in a concise, professional manner for"

def interact_with_bot1(user_input, defaults, context):
    # Hanya berinteraksi jika prompt terkait email
    email_keywords = ['write email', 'compose email', 'create email', 'email content']
    contains_email_keyword = any(keyword in user_input.lower() for keyword in email_keywords)

    if contains_email_keyword:
        response = palm.chat(
            **defaults,
            context=context,
            messages=[user_input]
        )
        return response.last
    else:
        return "The bot can only assist in email creation. Please enter the email related prompt."

# Fungsi untuk berinteraksi dengan AI untuk prompt yang terkait dengan CV
def interact_with_ai(user_input, defaults, context):
    cv_keywords = ['write cv', 'compose cv', 'create cv']
    contains_cv_keyword = any(keyword in user_input.lower() for keyword in cv_keywords)

    if contains_cv_keyword:
        response = palm.chat(
            **defaults,
            context=context,
            messages=[user_input]
        )
        return response.last
    else:
        return "The bot can only assist in CV creation. Please enter the CV related prompt."


# Fungsi untuk translate 10 bahasa UNESCO
def translate_to_indonesian(text):
    translator = Translator()
    translation = translator.translate(text, dest='id')
    return translation.text

def translate_to_english(text):
    translator = Translator()
    translation = translator.translate(text, dest='en')
    return translation.text

def translate_to_spanish(text):
    translator = Translator()
    translation = translator.translate(text, dest='es')
    return translation.text

def translate_to_french(text):
    translator = Translator()
    translation = translator.translate(text, dest='fr')
    return translation.text

def translate_to_hindi(text):
    translator = Translator()
    translation = translator.translate(text, dest='hi')
    return translation.text

def translate_to_russian(text):
    translator = Translator()
    translation = translator.translate(text, dest='ru')
    return translation.text

def translate_to_italian(text):
    translator = Translator()
    translation = translator.translate(text, dest='it')
    return translation.text

def translate_to_portuguese(text):
    translator = Translator()
    translation = translator.translate(text, dest='pt')
    return translation.text

def translate_to_arabic(text):
    translator = Translator()
    translation = translator.translate(text, dest='ar')
    return translation.text

def translate_to_mandarin(text):
    translator = Translator()
    translation = translator.translate(text, dest='zh-CN')
    return translation.text

def translate_response(response, language):
    if language == "Indonesian":
        return translate_to_indonesian(response)
    elif language == "English":
        return translate_to_english(response)
    elif language == "Spanish":
        return translate_to_spanish(response)
    elif language == "French":
        return translate_to_french(response)
    elif language == "Hindi":
        return translate_to_hindi(response)
    elif language == "Russian":
        return translate_to_russian(response)
    elif language == "Italian":
        return translate_to_italian(response)
    elif language == "Portuguese":
        return translate_to_portuguese(response)
    elif language == "Arabic":
        return translate_to_arabic(response)
    elif language == "Mandarin":
        return translate_to_mandarin(response)
    else:
        return response


# Fungsi untuk menyimpan ke file Word
def save_to_word_bytesio(content):
    doc_stream = BytesIO()
    doc = Document()
    doc.add_paragraph(content)
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# Fungsi untuk menyimpan ke file Word
def save_to_word(content, bot_option):
    # Mendapatkan timestamp saat ini untuk membuat nama file unik
    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")

    # Menghasilkan random string sebagai bagian dari nama file
    random_string = ''.join(random.choices(string.ascii_letters + string.digits, k=4))

    # Menambahkan data ke dokumen Word dengan nama file yang unik
    file_prefix = "E-Mail" if bot_option == "Create E-Mail" else "CV"
    file_name = f"{file_prefix}_{timestamp}_{random_string}.docx"
    doc = Document()
    doc.add_paragraph(content)
    doc.save(file_name)

    return file_name

st.title("Hunian Brutal Department")

bot_option = st.sidebar.radio("Select Bot:", ("Create E-Mail", "Create CV"))
if bot_option == "Create E-Mail":
    st.title("Create E- Mail")
    st.text("Gunakan keyword yang sesuai untuk membuat e-mail : ")
    st.text("'write email', 'compose email', 'create email', 'email content'")
    st.text(" ")
else:
    st.title("Create CV")
    st.text("Gunakan keyword yang sesuai untuk membuat cv : ")
    st.text("'write cv', 'compose cv', 'create cv'")
    st.text(" ")

st.sidebar.markdown(
    '<div style="position: fixed; bottom: 0; left: 0; width: 100%; padding: 10px; font-size: 15px; color: #ffff; box-shadow: 0px -1px 5px rgba(0, 0, 0, 0.1); display: flex; align-items: center; justify-content: flex-start;"><img src="https://cdn-icons-png.flaticon.com/512/106/106852.png" style="width: 20px; margin-right: 5px;"> Copyright by Hunian Brutal</div>',
    unsafe_allow_html=True
)

user_input = st.text_input("Prompt : ")
language_choice = st.selectbox("Select Language for Translation:", ("Indonesian", "English", "Spanish", "French", "Hindi", "Russian", "Italian", "Portuguese", "Arabic", "Mandarin"))
submit_button = st.button("Submit")
download_button = False

if submit_button:
    if user_input.strip() == "done":
        st.warning("Terima kasih! Anda telah menyelesaikan percakapan.")
    else:
        if bot_option == "Create E-Mail":
            ai_response = interact_with_bot1(user_input, defaults_bot1, context_bot1)
            download_button = True  # Mengaktifkan tombol unduh jika respons dari Bot 1 tersedia
        else:
            ai_response = interact_with_ai(user_input, defaults_bot1, context_bot2)
            download_button = True 
        
        translated_response = translate_response(ai_response, language_choice)  # Terjemahkan respon bot ke bahasa yang dipilih
        st.text_area("Result (Translated):", value=translated_response, height=200)

        if "The bot can only assist in email creation. Please enter the email related prompt." not in ai_response:
            if download_button:  # Menampilkan tombol unduh jika respons tersedia
                file_name = save_to_word(translated_response, bot_option)  # Simpan respon terjemahan ke dalam dokumen Word

                # Menampilkan tombol unduh dengan tautan ke file output.docx
                st.markdown(
                    f'<a href="data:file/docx;base64,{base64.b64encode(open(file_name, "rb").read()).decode()}" '
                    f'download="{file_name}" '
                    f'style="background-color:#008CBA;color:white;padding:10px 20px;text-decoration:none;border-radius:5px;display:inline-block;margin-top:10px;">'
                    f'Download Result as Word'
                    f'</a>',
                    unsafe_allow_html=True
                )
